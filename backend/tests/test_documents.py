"""Text extraction from uploaded solution documents.

This is the primary input path — the SoW is what most of the rubric is evaluated
against — and it had no coverage at all. PDFs and DOCX files are built for real
here rather than mocked: the point is to exercise pypdf and python-docx, so
stubbing them would test nothing.

Two defects these tests originally pinned are now fixed, and the tests assert the
fixed behaviour:

* every library failure is wrapped in `UnsupportedDocument` with a user-facing
  message, because the pipeline writes `type(exc).__name__` and the message into
  the status the UI displays;
* an extraction that yields no text raises instead of returning "", so a scanned
  SoW cannot be silently scored against nothing.
"""

from __future__ import annotations

import io

import pytest

from ingestion import documents


# --------------------------------------------------------------------------- #
# Fixture builders — real files, not mocks
# --------------------------------------------------------------------------- #

def make_pdf(pages: list[str]) -> bytes:
    """A real multi-page PDF with selectable text on the given pages."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    for body in pages:
        if body:
            pdf.drawString(72, 800, body)
        pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def make_scanned_pdf() -> bytes:
    """A PDF whose only content is a raster image — no text layer, as from a scanner."""
    from PIL import Image
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    image = io.BytesIO()
    Image.new("RGB", (600, 300), (210, 210, 210)).save(image, format="PNG")
    image.seek(0)

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    pdf.drawImage(ImageReader(image), 72, 500, 400, 200)
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def make_docx(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    from docx import Document

    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        added = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for col_index, cell in enumerate(row):
                added.cell(row_index, col_index).text = cell

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# PDF — the normal path
# --------------------------------------------------------------------------- #

def test_pdf_text_is_extracted_with_page_markers() -> None:
    data = make_pdf(["Statement of Work", "Orders are stored in DynamoDB"])

    text = documents.extract_text(data, "sow.pdf")

    assert "Statement of Work" in text
    assert "Orders are stored in DynamoDB" in text
    # Page markers give the model somewhere to anchor evidence quotes.
    assert "[page 1]" in text
    assert "[page 2]" in text
    assert text.index("[page 1]") < text.index("[page 2]")


def test_pdf_pages_with_no_text_are_omitted_rather_than_left_blank() -> None:
    data = make_pdf(["First page", "", "Third page"])

    text = documents.extract_text(data, "sow.pdf")

    assert "[page 1]" in text
    assert "[page 3]" in text
    # The empty page 2 contributes no marker, so numbering reflects real content.
    assert "[page 2]" not in text


def test_pdf_extension_matching_is_case_insensitive() -> None:
    data = make_pdf(["Uppercase extension"])

    assert "Uppercase extension" in documents.extract_text(data, "SOW.PDF")


# --------------------------------------------------------------------------- #
# PDF — malformed, empty, and scanned
# --------------------------------------------------------------------------- #

def test_an_empty_file_uploaded_as_a_pdf_raises_a_readable_error() -> None:
    with pytest.raises(documents.UnsupportedDocument, match="could not be read"):
        documents.extract_text(b"", "sow.pdf")


def test_a_corrupt_pdf_raises_rather_than_returning_partial_text() -> None:
    """Silently returning "" here would score a design against nothing."""
    with pytest.raises(documents.UnsupportedDocument, match="corrupt"):
        documents.extract_text(b"%PDF-1.4\nnot really a pdf", "sow.pdf")


def test_a_truncated_pdf_raises() -> None:
    whole = make_pdf(["Statement of Work", "Second page"])

    with pytest.raises(documents.UnsupportedDocument, match="truncated"):
        documents.extract_text(whole[: len(whole) // 2], "sow.pdf")


def test_a_non_pdf_renamed_to_pdf_raises() -> None:
    with pytest.raises(documents.UnsupportedDocument, match="not a PDF"):
        documents.extract_text(b"just plain text, not a pdf", "sow.pdf")


# The messages the UI will show verbatim. Kept in one place so a test failure
# points at the wording rather than at a scattered string literal.
LIBRARY_LEAKS = ("PdfStreamError", "EmptyFileError", "BadZipFile", "PdfReadError",
                 "Stream has ended", "not a zip file", "EOF marker", "pypdf",
                 "zipfile", "Traceback")


@pytest.mark.parametrize(
    ("data", "filename"),
    [
        (b"", "sow.pdf"),
        (b"%PDF-1.4 broken", "sow.pdf"),
        (b"just plain text", "sow.pdf"),
        (b"", "sow.docx"),
        (b"PK\x03\x04 garbage", "sow.docx"),
    ],
)
def test_library_failures_are_wrapped_in_a_catchable_domain_error(data, filename) -> None:
    """Was a pinned defect: pypdf and zipfile exceptions used to escape as-is.

    They must arrive as UnsupportedDocument so a caller can catch them as
    ValueError alongside the unsupported-extension case.
    """
    with pytest.raises(documents.UnsupportedDocument) as caught:
        documents.extract_text(data, filename)

    assert isinstance(caught.value, ValueError)


@pytest.mark.parametrize(
    ("data", "filename"),
    [
        (b"", "sow.pdf"),
        (b"%PDF-1.4 broken", "sow.pdf"),
        (b"just plain text", "sow.pdf"),
        (b"", "sow.docx"),
        (b"PK\x03\x04 garbage", "sow.docx"),
    ],
)
def test_no_library_internals_reach_the_user_facing_message(data, filename) -> None:
    """pipeline.run puts `type(exc).__name__: exc` straight into the status the UI
    renders, so the class name and the message both have to be presentable."""
    with pytest.raises(documents.UnsupportedDocument) as caught:
        documents.extract_text(data, filename)

    surfaced = f"{type(caught.value).__name__}: {caught.value}"
    for leak in LIBRARY_LEAKS:
        assert leak not in surfaced, f"{leak!r} leaked into: {surfaced}"


def test_the_cause_is_preserved_for_the_logs() -> None:
    """Wrapping is for the UI; the original traceback still has to be diagnosable."""
    with pytest.raises(documents.UnsupportedDocument) as caught:
        documents.extract_text(b"", "sow.pdf")

    assert caught.value.__cause__ is not None
    assert type(caught.value.__cause__).__module__.startswith("pypdf")


def test_a_password_protected_pdf_says_so_specifically() -> None:
    """Distinct from "corrupt": the user can act on this one."""
    from pypdf import PdfReader, PdfWriter

    writer = PdfWriter()
    writer.append(PdfReader(io.BytesIO(make_pdf(["Confidential SoW"]))))
    writer.encrypt("secret")
    encrypted = io.BytesIO()
    writer.write(encrypted)

    with pytest.raises(documents.UnsupportedDocument, match="password-protected"):
        documents.extract_text(encrypted.getvalue(), "sow.pdf")


def test_a_scanned_image_only_pdf_is_rejected_with_a_clear_message() -> None:
    """Was a pinned defect: this used to return "" silently.

    normalize.ingest only rejects an upload when the document *and* the diagram
    are both empty, so with a diagram attached the review used to proceed and be
    scored against a document nothing ever read. There is still no OCR; the point
    is that the user is told rather than misled.
    """
    with pytest.raises(documents.UnsupportedDocument) as caught:
        documents.extract_text(make_scanned_pdf(), "scanned-sow.pdf")

    message = str(caught.value)
    assert "No selectable text found" in message
    assert "scanned image" in message
    assert "OCR is not currently supported" in message


def test_a_pdf_whose_pages_are_all_blank_is_rejected_too() -> None:
    """Same failure mode without an image: nothing to review, so do not pretend."""
    with pytest.raises(documents.UnsupportedDocument, match="No selectable text"):
        documents.extract_text(make_pdf(["", "", ""]), "blank.pdf")


def test_a_docx_with_no_text_is_rejected_rather_than_returning_empty() -> None:
    """The .docx analogue of the scanned PDF — an images-only Word document."""
    with pytest.raises(documents.UnsupportedDocument, match="No text found"):
        documents.extract_text(make_docx([]), "images-only.docx")


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #

def test_docx_paragraphs_are_extracted() -> None:
    data = make_docx(["Statement of Work", "Scope and deliverables"])

    text = documents.extract_text(data, "sow.docx")

    assert "Statement of Work" in text
    assert "Scope and deliverables" in text


def test_docx_blank_paragraphs_are_dropped() -> None:
    data = make_docx(["First", "", "   ", "Second"])

    lines = documents.extract_text(data, "sow.docx").splitlines()

    assert lines == ["First", "Second"]


def test_docx_table_rows_are_flattened_into_pipe_separated_lines() -> None:
    """A SoW's requirements are often tabular, so losing tables loses the rubric's
    best evidence."""
    data = make_docx(
        ["Requirements"],
        table=[["Control", "Status"], ["Encryption at rest", "Not specified"]],
    )

    text = documents.extract_text(data, "sow.docx")

    assert "Control | Status" in text
    assert "Encryption at rest | Not specified" in text


def test_an_empty_file_uploaded_as_a_docx_raises_a_readable_error() -> None:
    with pytest.raises(documents.UnsupportedDocument, match="could not be read"):
        documents.extract_text(b"", "sow.docx")


def test_a_corrupt_docx_raises() -> None:
    """A .docx is a zip; a broken one fails at the zip layer, but the caller must
    not have to know that."""
    with pytest.raises(documents.UnsupportedDocument, match="Re-save it as .docx"):
        documents.extract_text(b"PK\x03\x04 truncated garbage", "sow.docx")


# --------------------------------------------------------------------------- #
# Plain-text formats
# --------------------------------------------------------------------------- #

@pytest.mark.parametrize(
    "filename",
    ["sow.txt", "sow.md", "sow.markdown", "sow.rst", "sow.csv", "sow.json",
     "sow.yaml", "sow.yml"],
)
def test_text_formats_are_decoded_as_utf8(filename: str) -> None:
    assert documents.extract_text(b"Design notes", filename) == "Design notes"


def test_surrounding_whitespace_is_stripped() -> None:
    assert documents.extract_text(b"\n\n  Design notes  \n\n", "sow.txt") == "Design notes"


def test_an_empty_text_file_yields_an_empty_string_without_raising() -> None:
    assert documents.extract_text(b"", "sow.txt") == ""


def test_invalid_utf8_is_replaced_rather_than_raising() -> None:
    """A Windows-encoded SoW must not take the whole review down."""
    text = documents.extract_text(b"Costs are \xa310,000 per month", "sow.txt")

    assert "Costs are" in text
    assert "10,000 per month" in text
    assert "�" in text, "the undecodable byte should become U+FFFD"


def test_unicode_survives_the_round_trip() -> None:
    assert "₹4,50,000" in documents.extract_text("Budget ₹4,50,000".encode(), "sow.md")


# --------------------------------------------------------------------------- #
# Dispatch and limits
# --------------------------------------------------------------------------- #

def test_an_unsupported_extension_raises_a_readable_domain_error() -> None:
    with pytest.raises(documents.UnsupportedDocument) as caught:
        documents.extract_text(b"anything", "sow.doc")

    message = str(caught.value)
    assert "sow.doc" in message
    # The legacy .doc case is common enough to warrant naming the fix.
    assert ".docx" in message


def test_a_file_with_no_extension_raises() -> None:
    with pytest.raises(documents.UnsupportedDocument):
        documents.extract_text(b"anything", "sow")


def test_an_oversized_document_is_truncated_with_a_visible_marker() -> None:
    """Guards the prompt against a pathological upload, and says so in the text so
    the model does not treat a cut-off document as complete."""
    oversized = ("x" * (documents.MAX_CHARS + 5_000)).encode()

    text = documents.extract_text(oversized, "sow.txt")

    assert len(text) < len(oversized.decode())
    assert text.endswith("[document truncated at the ingestion limit]")
    assert text.startswith("x" * 100)


def test_a_document_at_the_limit_is_not_truncated() -> None:
    exact = ("x" * documents.MAX_CHARS).encode()

    text = documents.extract_text(exact, "sow.txt")

    assert len(text) == documents.MAX_CHARS
    assert "truncated" not in text
