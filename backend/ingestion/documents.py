"""Text extraction from solution documents / SoWs.

Deterministic: PDFs via pypdf, .docx via python-docx, everything else decoded as
text. No model call.

Every failure leaves this module as `UnsupportedDocument`, which is a `ValueError`
and carries a message written for the person who uploaded the file. That matters
because `agent/pipeline.py` records `f"{type(exc).__name__}: {exc}"` into the
status the UI displays, so a raw `PdfStreamError: Stream has ended unexpectedly`
or `BadZipFile: File is not a zip file` would be shown to a user as-is.

An extraction that yields no text is also a failure, not an empty success. A
scanned SoW is a perfectly ordinary thing to be sent, and `normalize.ingest` only
rejects an upload when the document *and* the diagram are both empty — so
returning "" here would let a review run to completion and score a design against
a document nothing ever read.
"""

from __future__ import annotations

import io

# Guards the prompt against a pathological upload. Generous enough for a full SoW.
MAX_CHARS = 400_000

_TEXT_SUFFIXES = (
    ".txt", ".md", ".markdown", ".rst", ".csv", ".json", ".yaml", ".yml",
)


class UnsupportedDocument(ValueError):
    """The upload cannot be turned into reviewable text.

    The message is user-facing: it is surfaced in the UI verbatim.
    """


_PDF_NO_TEXT = (
    "No selectable text found — this looks like a scanned image. "
    "OCR is not currently supported."
)
_PDF_UNREADABLE = (
    "This PDF could not be read — it may be corrupt, truncated, or not a PDF "
    "file. Try opening it and re-exporting it as a PDF."
)
_PDF_ENCRYPTED = (
    "This PDF is password-protected, so its text cannot be read. Remove the "
    "protection and upload it again."
)
_DOCX_UNREADABLE = (
    "This Word document could not be read — it may be corrupt, or saved in the "
    "older .doc format. Re-save it as .docx and upload it again."
)
_DOCX_NO_TEXT = (
    "No text found in this Word document — it may contain only images. "
    "OCR is not currently supported."
)


def page_count(data: bytes, filename: str) -> int:
    """Total pages in a PDF, including pages that yield no text. 0 for anything else.

    Separate from `extract_text` because the extracted text CANNOT answer this. The
    `[page N]` markers `_pdf_text` writes are emitted only for pages that produced
    text, so the highest marker is the last READABLE page, not the last page — and a
    40-page document whose only text is on page 1 looks, from the text alone, exactly
    like a 1-page document.

    That distinction is the entire signal `quality.document_extraction` needs: the
    gap between "40 pages" and "1 page produced text" is what identifies a scanned
    upload. Reading it here costs one more `PdfReader` over bytes already in memory,
    with no page text extracted, and it never raises — `extract_text` runs first and
    owns every user-facing failure, so a count that cannot be taken is reported as
    0 (do not judge) rather than as an error.
    """
    if not filename.lower().endswith(".pdf"):
        return 0

    try:
        from pypdf import PdfReader

        return len(PdfReader(io.BytesIO(data)).pages)
    except Exception:  # noqa: BLE001 — a count, never a failure. See above.
        return 0


def extract_text(data: bytes, filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = _pdf_text(data)
    elif lower.endswith(".docx"):
        text = _docx_text(data)
    elif lower.endswith(_TEXT_SUFFIXES):
        # A plain-text file is allowed to be empty: unlike a PDF or a .docx, an
        # empty one means the file is empty, not that extraction failed.
        text = data.decode("utf-8", errors="replace")
    else:
        raise UnsupportedDocument(
            f"Cannot extract text from {filename!r}. Supported: .pdf, .docx, .txt, "
            ".md, .rst, .csv, .json, .yaml. (.doc is not supported — save as .docx.)"
        )

    text = text.strip()
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[document truncated at the ingestion limit]"
    return text


def _docx_text(data: bytes) -> str:
    """Paragraphs and table cells — a SoW's requirements are often tabular."""
    import zipfile

    from docx import Document
    from docx.opc.exceptions import OpcError

    try:
        document = Document(io.BytesIO(data))
    except (zipfile.BadZipFile, OpcError, KeyError) as exc:
        # BadZipFile: not a zip at all, or truncated. KeyError: a valid zip that
        # is not an Office package. Both are "we cannot read this" to a user.
        raise UnsupportedDocument(_DOCX_UNREADABLE) from exc

    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    if not text.strip():
        raise UnsupportedDocument(_DOCX_NO_TEXT)
    return text


def _pdf_text(data: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import DependencyError, PyPdfError

    try:
        reader = PdfReader(io.BytesIO(data))
    except PyPdfError as exc:
        # PyPdfError is the base of EmptyFileError, PdfStreamError and
        # PdfReadError, which is every way a malformed upload fails here.
        raise UnsupportedDocument(_PDF_UNREADABLE) from exc

    if reader.is_encrypted:
        raise UnsupportedDocument(_PDF_ENCRYPTED)

    try:
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except (PyPdfError, DependencyError) as exc:
        # DependencyError sits outside PyPdfError and is raised when a crypto
        # backend is missing, so it needs naming separately.
        raise UnsupportedDocument(_PDF_UNREADABLE) from exc

    text = "\n\n".join(f"[page {i}]\n{p}" for i, p in enumerate(pages, 1) if p)
    if not text.strip():
        raise UnsupportedDocument(_PDF_NO_TEXT)
    return text
